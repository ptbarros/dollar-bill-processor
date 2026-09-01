"""
Main Window - Dollar Bill Processor GUI
The central window containing all GUI components.
"""

import sys
from pathlib import Path
from typing import Optional

import cv2
import numpy as np

from PySide6.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout,
    QSplitter, QStatusBar, QMenuBar, QMenu, QFileDialog, QMessageBox,
    QProgressBar, QLabel, QPushButton
)
from PySide6.QtCore import Qt, Signal, Slot, QThread, QTimer
from PySide6.QtGui import QAction, QKeySequence, QShortcut, QPixmap, QImage

# Import our components
from .processing_panel import ProcessingPanel
from .results_list import ResultsList
from .preview_panel import PreviewPanel
from .layout_manager import (
    LayoutManager, LAYOUT_CLASSIC, LAYOUT_WIDE_PREVIEW,
    LAYOUT_DETAILS_RIGHT, LAYOUT_NAMES
)

# Import backend
sys.path.insert(0, str(Path(__file__).parent.parent))
from settings_manager import SettingsManager, get_settings
from version import get_version_string
from correction_manager import CorrectionManager
from session_recovery import SessionRecoveryManager, get_recovery_manager
from debug_logger import dlog, fingerprint


class MainWindow(QMainWindow):
    """Main application window."""

    def __init__(self):
        super().__init__()

        # Load settings
        self.settings = get_settings()
        self.correction_manager = CorrectionManager()
        self.recovery_manager = get_recovery_manager()

        # Processing state
        self.processor = None
        self.current_results = []
        self.is_processing = False

        # Monitor mode state
        self.is_monitoring = False
        self.file_watcher = None
        self.monitor_thread = None

        # Autosave state
        self._session_dirty = False
        self._current_input_dir = ""

        # Debounce for alignment warnings (prevent spam during batch load)
        self._last_align_warning_time = 0

        # Setup UI
        self._setup_ui()
        self._setup_menus()
        self._setup_shortcuts()
        self._setup_statusbar()
        self._restore_geometry()
        self._setup_autosave_timer()
        self._apply_settings()  # Apply font size and other settings

        self.setWindowTitle("Dollar Bill Processor")

        # Check for recovery on startup (after UI is ready)
        QTimer.singleShot(100, self._check_for_recovery)

    def _setup_ui(self):
        """Setup the main UI layout."""
        central_widget = QWidget()
        self.setCentralWidget(central_widget)

        self.main_layout = QVBoxLayout(central_widget)
        self.main_layout.setContentsMargins(8, 8, 8, 8)
        self.main_layout.setSpacing(8)

        # Top toolbar area
        self.processing_panel = ProcessingPanel()
        self.processing_panel.process_requested.connect(self._on_process_requested)
        self.processing_panel.organize_requested.connect(self._on_organize_requested)
        self.processing_panel.stop_requested.connect(self._on_stop_requested)
        self.processing_panel.monitor_requested.connect(self._start_monitoring)
        self.processing_panel.monitor_stop_requested.connect(self._stop_monitoring)
        self.processing_panel.monitor_check.toggled.connect(self._on_monitor_mode_changed)
        self.processing_panel.archive_requested.connect(self._on_archive_requested)
        self.main_layout.addWidget(self.processing_panel)

        # Create panels (not added to layout yet - LayoutManager will do that)
        self.results_list = ResultsList()
        self.results_list.item_selected.connect(self._on_result_selected)
        self.results_list.correction_applied.connect(self._on_correction_applied)
        self.results_list.batch_changed.connect(self._on_batch_changed)
        self.results_list.status_changed.connect(self._mark_session_dirty)

        self.preview_panel = PreviewPanel()
        self.preview_panel.prev_requested.connect(self._prev_bill)
        self.preview_panel.next_requested.connect(self._next_bill)
        self.preview_panel.align_requested.connect(self._on_align_image)
        self.preview_panel.px_dev_updated.connect(self.results_list.update_px_dev)
        self.preview_panel.crop_requested.connect(self._on_crop_current)
        self.results_list.crop_requested.connect(self._on_crop_selected)
        # Apply saved visibility settings
        self.preview_panel.set_serial_region_visible(self.settings.ui.show_serial_region)
        self.preview_panel.set_details_visible(self.settings.ui.show_bill_details)

        # Initialize layout manager and apply saved layout
        self.layout_manager = LayoutManager(central_widget)
        self.layout_manager.set_widgets(self.results_list, self.preview_panel, self.processing_panel)
        self.layout_manager.set_parent_layout(self.main_layout)

        # Apply the saved layout (or default to classic)
        saved_layout = self.settings.ui.layout_mode
        if saved_layout not in (LAYOUT_CLASSIC, LAYOUT_WIDE_PREVIEW, LAYOUT_DETAILS_RIGHT):
            saved_layout = LAYOUT_CLASSIC
        self.layout_manager.apply_layout(saved_layout)

    def _setup_menus(self):
        """Setup the menu bar."""
        menubar = self.menuBar()

        # File menu
        file_menu = menubar.addMenu("&File")

        open_action = QAction("&Open Folder...", self)
        open_action.setShortcut(QKeySequence.Open)
        open_action.triggered.connect(self._on_open_folder)
        file_menu.addAction(open_action)

        file_menu.addSeparator()

        export_menu = file_menu.addMenu("&Export Results")

        export_csv = QAction("Export as CSV...", self)
        export_csv.triggered.connect(lambda: self._export_results("csv"))
        export_menu.addAction(export_csv)

        export_excel = QAction("Export as Excel...", self)
        export_excel.triggered.connect(lambda: self._export_results("excel"))
        export_menu.addAction(export_excel)

        export_html = QAction("Export as HTML Report...", self)
        export_html.triggered.connect(lambda: self._export_results("html"))
        export_menu.addAction(export_html)

        file_menu.addSeparator()

        exit_action = QAction("E&xit", self)
        exit_action.setShortcut(QKeySequence.Quit)
        exit_action.triggered.connect(self.close)
        file_menu.addAction(exit_action)

        # Edit menu
        edit_menu = menubar.addMenu("&Edit")

        settings_action = QAction("&Settings...", self)
        settings_action.setShortcut(QKeySequence("Ctrl+,"))
        settings_action.triggered.connect(self._on_settings)
        edit_menu.addAction(settings_action)

        edit_menu.addSeparator()

        patterns_action = QAction("&Pattern Manager...", self)
        patterns_action.triggered.connect(self._on_pattern_manager)
        edit_menu.addAction(patterns_action)

        crops_action = QAction("eBay &Crop Manager...", self)
        crops_action.triggered.connect(self._on_crop_manager)
        edit_menu.addAction(crops_action)

        # View menu
        view_menu = menubar.addMenu("&View")

        review_only = QAction("Show &Review Items Only", self, checkable=True)
        review_only.triggered.connect(self._toggle_review_filter)
        view_menu.addAction(review_only)

        fancy_only = QAction("Show &Fancy Bills Only", self, checkable=True)
        fancy_only.triggered.connect(self._toggle_fancy_filter)
        view_menu.addAction(fancy_only)

        view_menu.addSeparator()

        # Panel visibility toggles (load from settings)
        self.show_serial_region_action = QAction("Show &Serial Region", self, checkable=True)
        self.show_serial_region_action.setChecked(self.settings.ui.show_serial_region)
        self.show_serial_region_action.triggered.connect(self._toggle_serial_region)
        view_menu.addAction(self.show_serial_region_action)

        self.show_details_action = QAction("Show Bill &Details", self, checkable=True)
        self.show_details_action.setChecked(self.settings.ui.show_bill_details)
        self.show_details_action.triggered.connect(self._toggle_details)
        view_menu.addAction(self.show_details_action)

        view_menu.addSeparator()

        # Layout submenu
        layout_menu = view_menu.addMenu("&Layout")
        self.layout_actions = {}
        current_layout = self.settings.ui.layout_mode

        for layout_id, layout_name in LAYOUT_NAMES.items():
            action = QAction(layout_name, self, checkable=True)
            action.setChecked(layout_id == current_layout)
            action.triggered.connect(lambda checked, lid=layout_id: self._on_layout_changed(lid))
            layout_menu.addAction(action)
            self.layout_actions[layout_id] = action

        # Columns submenu - toggle visibility of table columns
        columns_menu = view_menu.addMenu("&Columns")
        self.column_actions = {}
        self._setup_columns_menu(columns_menu)

        view_menu.addSeparator()

        refresh_action = QAction("&Refresh", self)
        refresh_action.setShortcut(QKeySequence.Refresh)
        refresh_action.triggered.connect(self._refresh_view)
        view_menu.addAction(refresh_action)

        # Connect callback to sync column menu when hiding via header right-click
        self.results_list.set_column_visibility_callback(self._on_column_hidden_from_header)

        # Help menu
        help_menu = menubar.addMenu("&Help")

        about_action = QAction("&About...", self)
        about_action.triggered.connect(self._on_about)
        help_menu.addAction(about_action)

    def _setup_shortcuts(self):
        """Setup keyboard shortcuts for navigation and zoom."""
        # Bill navigation
        QShortcut(QKeySequence(Qt.Key_PageDown), self, self._next_bill)
        QShortcut(QKeySequence(Qt.Key_PageUp), self, self._prev_bill)
        QShortcut(QKeySequence(Qt.Key_N), self, self._next_bill)
        QShortcut(QKeySequence(Qt.Key_P), self, self._prev_bill)
        QShortcut(QKeySequence(Qt.Key_Down), self, self._next_bill)
        QShortcut(QKeySequence(Qt.Key_Up), self, self._prev_bill)

        # Crop shortcut - batch crop all queued (checked) bills
        QShortcut(QKeySequence(Qt.Key_C), self, self._on_batch_crop_queued)

        # Zoom controls
        QShortcut(QKeySequence(Qt.Key_Plus), self, self._zoom_in)
        QShortcut(QKeySequence(Qt.Key_Equal), self, self._zoom_in)  # = key (unshifted +)
        QShortcut(QKeySequence(Qt.Key_Minus), self, self._zoom_out)
        QShortcut(QKeySequence(Qt.Key_0), self, self._zoom_fit)
        QShortcut(QKeySequence(Qt.Key_F), self, self._zoom_fit)

        # Pan controls (Shift + arrow keys)
        QShortcut(QKeySequence(Qt.SHIFT | Qt.Key_Left), self, self._pan_left)
        QShortcut(QKeySequence(Qt.SHIFT | Qt.Key_Right), self, self._pan_right)
        QShortcut(QKeySequence(Qt.SHIFT | Qt.Key_Up), self, self._pan_up)
        QShortcut(QKeySequence(Qt.SHIFT | Qt.Key_Down), self, self._pan_down)

        # Toggle checked (review status)
        space_shortcut = QShortcut(QKeySequence(Qt.Key_Space), self, self._toggle_checked)
        space_shortcut.setContext(Qt.ApplicationShortcut)

        # Plate magnifier (for mule comparison)
        QShortcut(QKeySequence(Qt.Key_M), self, self._show_plate_magnifier)

        # Auto-Align toggle
        QShortcut(QKeySequence(Qt.Key_A), self, self._toggle_auto_align)

        # Crosshair toggle
        QShortcut(QKeySequence(Qt.Key_X), self, self._toggle_crosshair)

        # View mode shortcuts (1-5)
        QShortcut(QKeySequence(Qt.Key_1), self, lambda: self._switch_view("front"))
        QShortcut(QKeySequence(Qt.Key_2), self, lambda: self._switch_view("back"))
        QShortcut(QKeySequence(Qt.Key_3), self, lambda: self._switch_view("stitched"))
        QShortcut(QKeySequence(Qt.Key_4), self, lambda: self._switch_view("split_v"))
        QShortcut(QKeySequence(Qt.Key_5), self, lambda: self._switch_view("split_h"))

    def _next_bill(self):
        """Navigate to next bill in results."""
        current = self.results_list.tree.currentItem()
        if current:
            index = self.results_list.tree.indexOfTopLevelItem(current)
            if index < self.results_list.tree.topLevelItemCount() - 1:
                next_item = self.results_list.tree.topLevelItem(index + 1)
                self.results_list.tree.setCurrentItem(next_item)

    def _prev_bill(self):
        """Navigate to previous bill in results."""
        current = self.results_list.tree.currentItem()
        if current:
            index = self.results_list.tree.indexOfTopLevelItem(current)
            if index > 0:
                prev_item = self.results_list.tree.topLevelItem(index - 1)
                self.results_list.tree.setCurrentItem(prev_item)

    def _on_align_image(self, image_path: str):
        """Handle alignment request from preview panel."""
        # Check if showing aligned - if so, reset instead
        if self.preview_panel._is_showing_aligned:
            self.preview_panel.reset_aligned_image()
            return

        # Get processor - could be from batch processing or monitor mode
        processor = self.processor
        if not processor and hasattr(self, 'processing_thread') and self.processing_thread:
            processor = self.processing_thread.processor
        if not processor and hasattr(self, 'monitor_thread') and self.monitor_thread:
            processor = self.monitor_thread.processor

        # Check for cached alignment values from archived batch
        current_result = self.preview_panel.current_result
        cached_angle = current_result.get('front_align_angle', 0.0) if current_result else 0.0
        cached_flipped = current_result.get('front_align_flipped', False) if current_result else False
        # Check if alignment data was present in CSV (set during _load_batch)
        # This properly handles archives where angle was 0.0 (no rotation needed)
        has_cached_alignment = current_result.get('_has_alignment_data', False) if current_result else False

        if not processor and not has_cached_alignment:
            # Debounce: only show warning if 2+ seconds since last warning
            # (prevents spam when auto-align fires for multiple items during batch load)
            import time
            current_time = time.time()
            if current_time - self._last_align_warning_time >= 2.0:
                self._last_align_warning_time = current_time
                QMessageBox.warning(self, "No Alignment Data",
                    "No alignment data available.\n\n"
                    "For archived batches processed before this update, "
                    "rotation values weren't saved. Reprocess the folder to enable alignment.")
            return

        if not image_path:
            return

        try:
            # Align both front and back images
            front_path = self.preview_panel._current_front_file
            back_path = self.preview_panel._current_back_file

            front_pixmap = None
            back_pixmap = None
            status_msg = ""
            front_angle = 0.0
            front_flipped = False

            # Try cached alignment first, fall back to YOLO detection
            if front_path:
                if has_cached_alignment:
                    # Use cached values from archived batch (no YOLO needed)
                    aligned_img = self._apply_cached_alignment(Path(front_path), cached_angle, cached_flipped)
                    front_angle = cached_angle
                    front_flipped = cached_flipped
                    status_msg = f"Aligned (cached): {front_angle:.1f}° rotation"
                elif processor:
                    # Use YOLO detection
                    aligned_img, info = processor.align_for_preview(Path(front_path))
                    front_angle = info.get('angle', 0) if info else 0
                    front_flipped = info.get('flipped', False) if info else False
                    status_msg = f"Aligned: {front_angle:.1f}° rotation"
                else:
                    aligned_img = None

                if aligned_img is not None:
                    front_pixmap = self._cv2_to_pixmap(aligned_img)
                    if front_flipped:
                        status_msg += ", flipped 180°"

            # Align back using OPPOSITE rotation from the front
            # Physical flip mirrors the skew: if front is +2° CW, back appears -2° CCW
            if back_path and Path(back_path).exists():
                aligned_back = self._apply_cached_alignment(Path(back_path), -front_angle, front_flipped)
                if aligned_back is not None:
                    back_pixmap = self._cv2_to_pixmap(aligned_back)

            if front_pixmap is None and back_pixmap is None:
                QMessageBox.warning(self, "Alignment Failed", "Could not align the images.")
                return

            self.statusBar().showMessage(status_msg, 5000)

            # Display the aligned images in all views
            self.preview_panel.show_aligned_images(front_pixmap, back_pixmap)

        except Exception as e:
            QMessageBox.warning(self, "Alignment Error", f"Error during alignment: {str(e)}")

    def _cv2_to_pixmap(self, cv2_img) -> QPixmap:
        """Convert OpenCV BGR image to QPixmap."""
        h, w, ch = cv2_img.shape
        bytes_per_line = ch * w
        # Convert BGR to RGB for Qt
        rgb_img = cv2_img[:, :, ::-1].copy()
        q_img = QImage(rgb_img.data, w, h, bytes_per_line, QImage.Format_RGB888)
        return QPixmap.fromImage(q_img)

    def _apply_cached_alignment(self, image_path: Path, angle: float, flipped: bool) -> Optional[np.ndarray]:
        """Apply alignment using cached rotation values (no YOLO needed).

        This enables alignment on archived batches without reprocessing.
        """
        img = cv2.imread(str(image_path))
        if img is None:
            return None

        h, w = img.shape[:2]

        # Apply rotation if needed (0.8° threshold)
        if abs(angle) >= 0.8:
            center = (w // 2, h // 2)
            M = cv2.getRotationMatrix2D(center, angle, 1.0)
            img = cv2.warpAffine(img, M, (w, h),
                                 flags=cv2.INTER_CUBIC,
                                 borderMode=cv2.BORDER_CONSTANT,
                                 borderValue=(255, 255, 255))

        # Apply 180° flip if needed
        if flipped:
            img = cv2.rotate(img, cv2.ROTATE_180)

        return img

    def _on_batch_crop_queued(self):
        """Generate crops for all queued (checked) bills."""
        queued = [r for r in self.results_list.results if r.get('checked')]
        if not queued:
            self.statusBar().showMessage("No bills queued (Space to queue)", 3000)
            return
        self._on_crop_selected(queued)

    def _on_crop_current(self):
        """Generate crops for the currently displayed bill."""
        result = self.preview_panel.current_result
        if result:
            self._on_crop_selected([result])

    def _on_crop_selected(self, results: list):
        """Generate crops for selected bills."""
        dlog("action.crop_selected", count=len(results) if results else 0,
             viewing_batch=str(self.results_list.get_current_batch_path() or "(current)"),
             state=fingerprint(self.current_results))
        if not results:
            return

        # Get processor - try existing ones first
        processor = self.processor
        if not processor and hasattr(self, 'processing_thread') and self.processing_thread:
            processor = self.processing_thread.processor
        if not processor and hasattr(self, 'monitor_thread') and self.monitor_thread:
            processor = self.monitor_thread.processor

        # If no processor exists, create one lazily for cropping
        if not processor:
            processor = self._get_or_create_processor()
            if not processor:
                return  # User cancelled or error occurred

        # Get output directory - use last output dir or ask
        output_dir = Path(self.settings.ui.last_output_dir) if self.settings.ui.last_output_dir else None
        if not output_dir or not output_dir.exists():
            output_dir = QFileDialog.getExistingDirectory(
                self, "Select Output Directory for Crops",
                str(Path.home())
            )
            if not output_dir:
                return
            output_dir = Path(output_dir)

        # Reload config to pick up any changes from eBay Crop Manager
        processor.cfg.reload()

        # Generate crops for each result
        from process_production import BillPair
        cropped_count = 0

        for result in results:
            try:
                # Create a BillPair from the result
                pair = BillPair(
                    front_path=Path(result.get('front_file', '')),
                    back_path=Path(result.get('back_file', '')) if result.get('back_file') else None,
                    stack_position=result.get('position', 0),
                    serial=result.get('serial', ''),
                    confidence=float(result.get('confidence', 0)),
                )
                # Set alignment info for crop generation
                pair.front_align_angle = result.get('front_align_angle', 0.0)
                pair.front_align_flipped = result.get('front_align_flipped', False)
                # Pattern info drives the serial overlay crop(s). pattern_overrides
                # (right-click "Set Pattern(s)...") picks which overlays to draw --
                # one crop each; fancy_types is the fallback set of matched patterns.
                fancy_types_str = result.get('fancy_types', '') or ''
                pair.fancy_types = [p.strip() for p in fancy_types_str.split(',') if p.strip()]
                pair.pattern_override = result.get('pattern_override')
                overrides_str = result.get('pattern_overrides') or ''
                pair.pattern_overrides = [p.strip() for p in overrides_str.split(',') if p.strip()]

                # Generate crops
                processor.generate_crops(pair, output_dir)
                cropped_count += 1
                result['cropped'] = True

            except Exception as e:
                print(f"Error cropping {result.get('serial', 'unknown')}: {e}")

        # Update status cells for cropped results (mark_cropped also clears checked flag)
        cropped_results = [r for r in results if r.get('cropped')]
        if cropped_results:
            self.results_list.mark_cropped(cropped_results)
            self._mark_session_dirty()

        # Generate printable labels file
        self._generate_labels(results, output_dir)
        self._generate_labels_docx(results, output_dir)

        # Show confirmation
        self.statusBar().showMessage(f"Generated crops for {cropped_count} bill(s) in {output_dir}", 5000)

    def _pattern_display_name(self, pattern_str: str) -> str:
        """Resolve a pattern identifier to its friendly display name for labels.

        Detected patterns are stored as internal names (e.g. STEP_LADDER); the
        engine maps those to a friendly display name. A user-typed override that
        isn't a known pattern has no mapping and passes through unchanged.
        """
        if not pattern_str or pattern_str == 'No Pattern':
            return pattern_str
        try:
            return self.results_list._get_display_name(pattern_str)
        except Exception:
            return pattern_str

    def _selected_patterns(self, result: dict) -> list:
        """Resolve which pattern name(s) a bill's label/crops should use.

        Priority: multi-select pattern_overrides -> single pattern_override ->
        the first detected fancy type. Returns a list (may be empty).
        """
        overrides = result.get('pattern_overrides')
        if overrides:
            return [p.strip() for p in overrides.split(',') if p.strip()]
        single = result.get('pattern_override')
        if single:
            return [single]
        patterns = [p.strip() for p in (result.get('fancy_types', '') or '').split(',') if p.strip()]
        return [patterns[0]] if patterns else []

    def _generate_labels(self, results: list, output_dir: Path):
        """Generate printable labels file for bills.

        Each bill gets two label entries:
        1. Without catalog (for reference/binder)
        2. With catalog (to store with the physical bill)
        """
        labels_file = output_dir / "zbill_labels.txt"

        with open(labels_file, 'a', encoding='utf-8') as f:
            for result in results:
                serial = result.get('serial', 'Unknown')
                position = result.get('position', '')
                series = result.get('series_year', '')

                # Selected pattern(s): may be several (multi-select), each with
                # its own catalog number.
                selected = self._selected_patterns(result)
                if not selected:
                    selected = ['No Pattern']
                pattern_display = ', '.join(self._pattern_display_name(p) for p in selected)
                catalog = ', '.join(
                    c for c in (self.settings.get_pattern_catalog(p, '') for p in selected
                                if p != 'No Pattern') if c
                )

                # Get user note if present
                note = result.get('note', '')

                # Label 1: Without catalog (for reference)
                f.write(f"Serial: {serial}\n")
                f.write(f"Pattern: {pattern_display}\n")
                if note:
                    f.write(f"Note: {note}\n")
                f.write(f"Series: {series}\n")
                f.write("-" * 30 + "\n\n")

                # Label 2: With catalog (to store with bill)
                f.write(f"Serial: {serial}\n")
                f.write(f"Pattern: {pattern_display}\n")
                if note:
                    f.write(f"Note: {note}\n")
                f.write(f"Series: {series}\n")
                f.write(f"Catalog: {catalog}  Pos: {position}\n")
                f.write("=" * 30 + "\n\n")

    def _generate_labels_docx(self, results: list, output_dir: Path):
        """Generate printable .docx labels for Rollo label printer (2x1 inch labels).

        Each bill gets two label pages:
        1. Without catalog (for reference/binder)
        2. With catalog (to store with the physical bill)
        """
        try:
            from docx import Document
            from docx.shared import Inches, Pt, Twips
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.oxml.ns import qn
        except ImportError:
            print("python-docx not installed, skipping .docx label generation")
            return

        labels_file = output_dir / "zbill_labels.docx"
        doc = Document()

        first_label = True
        for result in results:
            serial = result.get('serial', 'Unknown')
            position = result.get('position', '')
            series = result.get('series_year', '')

            # Selected pattern(s): may be several (multi-select).
            selected = self._selected_patterns(result)
            if not selected:
                selected = ['No Pattern']
            pattern_display = ', '.join(self._pattern_display_name(p) for p in selected)
            catalog = ', '.join(
                c for c in (self.settings.get_pattern_catalog(p, '') for p in selected
                            if p != 'No Pattern') if c
            )

            note = result.get('note', '')

            for label_num in range(2):
                # Set up page as 2x1 inch label with zero margins
                if first_label:
                    section = doc.sections[0]
                    first_label = False
                else:
                    section = doc.add_section()

                section.page_width = Inches(2)
                section.page_height = Inches(1)
                section.left_margin = Inches(0)
                section.right_margin = Inches(0)
                section.top_margin = Inches(0)
                section.bottom_margin = Inches(0)

                # Create table (single cell, matching label.docx structure)
                table = doc.add_table(rows=1, cols=1)
                table.autofit = False

                # Remove table borders
                tbl_pr = table._element.find(qn('w:tblPr'))
                borders = tbl_pr.find(qn('w:tblBorders'))
                if borders is None:
                    borders = tbl_pr.makeelement(qn('w:tblBorders'), {})
                    tbl_pr.append(borders)
                for border_name in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
                    border_el = borders.makeelement(qn(f'w:{border_name}'), {
                        qn('w:val'): 'none', qn('w:sz'): '0',
                        qn('w:space'): '0', qn('w:color'): 'auto'
                    })
                    existing = borders.find(qn(f'w:{border_name}'))
                    if existing is not None:
                        borders.remove(existing)
                    borders.append(border_el)

                # Set fixed layout and cell width
                tbl_layout = tbl_pr.makeelement(qn('w:tblLayout'), {qn('w:type'): 'fixed'})
                tbl_pr.append(tbl_layout)

                # Set cell margins to minimal (matching label.docx: 15 twips L/R)
                cell_mar = tbl_pr.makeelement(qn('w:tblCellMar'), {})
                for side, val in [('left', '15'), ('right', '15')]:
                    el = cell_mar.makeelement(qn(f'w:{side}'), {
                        qn('w:w'): val, qn('w:type'): 'dxa'
                    })
                    cell_mar.append(el)
                tbl_pr.append(cell_mar)

                # Set column width to 2 inches (2880 twips)
                grid = table._element.find(qn('w:tblGrid'))
                if grid is not None:
                    for col in grid.findall(qn('w:gridCol')):
                        col.set(qn('w:w'), '2880')

                # Set row height to 1 inch exact (1440 twips)
                row = table.rows[0]
                tr_pr = row._element.find(qn('w:trPr'))
                if tr_pr is None:
                    tr_pr = row._element.makeelement(qn('w:trPr'), {})
                    row._element.insert(0, tr_pr)
                tr_height = tr_pr.makeelement(qn('w:trHeight'), {
                    qn('w:val'): '1440', qn('w:hRule'): 'exact'
                })
                tr_pr.append(tr_height)

                cell = table.cell(0, 0)
                # Set cell width
                tc_pr = cell._element.find(qn('w:tcPr'))
                if tc_pr is None:
                    tc_pr = cell._element.makeelement(qn('w:tcPr'), {})
                    cell._element.insert(0, tc_pr)
                tc_w = tc_pr.makeelement(qn('w:tcW'), {
                    qn('w:w'): '2880', qn('w:type'): 'dxa'
                })
                tc_pr.append(tc_w)

                # Line 1: serial + series year
                p1 = cell.paragraphs[0]
                p1_ppr = p1._element.find(qn('w:pPr'))
                if p1_ppr is None:
                    p1_ppr = p1._element.makeelement(qn('w:pPr'), {})
                    p1._element.insert(0, p1_ppr)
                # Add spacing before (matching label.docx: 111 twips)
                spacing = p1_ppr.makeelement(qn('w:spacing'), {qn('w:before'): '111'})
                p1_ppr.append(spacing)
                # Add indent (matching label.docx: 72 twips L/R)
                indent = p1_ppr.makeelement(qn('w:ind'), {
                    qn('w:left'): '72', qn('w:right'): '72'
                })
                p1_ppr.append(indent)

                run1 = p1.add_run(f"{serial}     SERIES  {series}")
                run1.font.size = Pt(10)

                # Line 2: pattern name
                p2 = cell.add_paragraph()
                p2_ppr = p2._element.find(qn('w:pPr'))
                if p2_ppr is None:
                    p2_ppr = p2._element.makeelement(qn('w:pPr'), {})
                    p2._element.insert(0, p2_ppr)
                indent2 = p2_ppr.makeelement(qn('w:ind'), {
                    qn('w:left'): '72', qn('w:right'): '72'
                })
                p2_ppr.append(indent2)

                run2 = p2.add_run(pattern_display)
                run2.font.size = Pt(10)

                # Line 3: note or catalog info
                line3 = ''
                if label_num == 0 and note:
                    line3 = note
                elif label_num == 1:
                    if note:
                        line3 = note
                    catalog_line = f"Catalog: {catalog}  Pos: {position}" if catalog or position else ''
                    if catalog_line:
                        line3 = f"{line3}\n{catalog_line}" if line3 else catalog_line

                if line3:
                    for li, line_text in enumerate(line3.split('\n')):
                        p3 = cell.add_paragraph()
                        p3_ppr = p3._element.find(qn('w:pPr'))
                        if p3_ppr is None:
                            p3_ppr = p3._element.makeelement(qn('w:pPr'), {})
                            p3._element.insert(0, p3_ppr)
                        indent3 = p3_ppr.makeelement(qn('w:ind'), {
                            qn('w:left'): '72', qn('w:right'): '72'
                        })
                        p3_ppr.append(indent3)
                        run3 = p3.add_run(line_text)
                        run3.font.size = Pt(10)

        # Remove the empty paragraph that python-docx adds at the start
        body = doc.element.body
        first_p = body.find(qn('w:p'))
        first_tbl = body.find(qn('w:tbl'))
        if first_p is not None and first_tbl is not None:
            if list(body).index(first_p) < list(body).index(first_tbl):
                body.remove(first_p)

        doc.save(str(labels_file))

    def _zoom_in(self):
        """Zoom in on preview."""
        self.preview_panel.zoom_in()

    def _zoom_out(self):
        """Zoom out on preview."""
        self.preview_panel.zoom_out()

    def _zoom_fit(self):
        """Fit zoom on preview."""
        self.preview_panel.zoom_fit()

    def _pan_left(self):
        """Pan preview left."""
        self.preview_panel.pan(-50, 0)

    def _pan_right(self):
        """Pan preview right."""
        self.preview_panel.pan(50, 0)

    def _pan_up(self):
        """Pan preview up."""
        self.preview_panel.pan(0, -50)

    def _pan_down(self):
        """Pan preview down."""
        self.preview_panel.pan(0, 50)

    def _toggle_checked(self):
        """Toggle checked status on currently selected bill(s)."""
        self.results_list.toggle_checked()
        self._mark_session_dirty()

    def _show_plate_magnifier(self):
        """Show plate magnifier popup for mule comparison."""
        self.preview_panel.show_plate_magnifier()

    def _toggle_auto_align(self):
        """Toggle auto-align via keyboard shortcut."""
        btn = self.preview_panel.align_btn
        btn.setChecked(not btn.isChecked())
        self.preview_panel._on_align_toggled(btn.isChecked())

    def _toggle_crosshair(self):
        """Toggle crosshair overlay via keyboard shortcut."""
        btn = self.preview_panel.crosshair_btn
        btn.setChecked(not btn.isChecked())
        self.preview_panel._on_crosshair_toggled(btn.isChecked())

    def _switch_view(self, mode: str):
        """Switch preview view mode via keyboard shortcut."""
        self.preview_panel._on_view_mode_clicked(mode)

    def _setup_statusbar(self):
        """Setup the status bar."""
        self.statusbar = QStatusBar()
        self.setStatusBar(self.statusbar)

        # Status label
        self.status_label = QLabel("Ready")
        self.statusbar.addWidget(self.status_label, 1)

        # Progress info
        self.progress_label = QLabel("")
        self.statusbar.addPermanentWidget(self.progress_label)

    def _restore_geometry(self):
        """Restore window geometry from settings."""
        x, y, w, h = self.settings.get_window_geometry()
        self.setGeometry(x, y, w, h)

    def _save_geometry(self):
        """Save window geometry to settings."""
        geo = self.geometry()
        self.settings.update_window_geometry(geo.x(), geo.y(), geo.width(), geo.height())
        self.settings.save()

    def closeEvent(self, event):
        """Handle window close."""
        self._save_geometry()
        if self.is_processing:
            reply = QMessageBox.question(
                self, "Confirm Exit",
                "Processing is in progress. Are you sure you want to exit?",
                QMessageBox.Yes | QMessageBox.No,
                QMessageBox.No
            )
            if reply == QMessageBox.No:
                event.ignore()
                return
            self._on_stop_requested()
        if self.is_monitoring:
            reply = QMessageBox.question(
                self, "Confirm Exit",
                "Monitoring is active. Are you sure you want to exit?\n\n"
                "Files will be archived if auto-archive is enabled.",
                QMessageBox.Yes | QMessageBox.No,
                QMessageBox.No
            )
            if reply == QMessageBox.No:
                event.ignore()
                return
            self._stop_monitoring()

        # Final autosave before exit (if we have unarchived results)
        if self.current_results and self._session_dirty:
            self._trigger_autosave()

        # Clean up resources to prevent memory corruption on exit
        self._cleanup_resources()

        event.accept()

    def _cleanup_resources(self):
        """Clean up native resources (YOLO, EasyOCR, threads) before exit.

        This prevents heap corruption caused by non-deterministic cleanup order
        between Qt, Python GC, and native extensions.
        """
        import gc

        # Stop autosave timer
        if hasattr(self, '_autosave_timer') and self._autosave_timer:
            self._autosave_timer.stop()

        # Stop file watcher first
        if self.file_watcher:
            self.file_watcher.stop()
            self.file_watcher.wait(1000)
            self.file_watcher = None

        # Wait for processing thread to finish
        if hasattr(self, 'processing_thread') and self.processing_thread:
            if self.processing_thread.isRunning():
                self.processing_thread.request_stop()
                self.processing_thread.wait(3000)  # Wait up to 3 seconds
            # Clear processor reference in thread before releasing thread
            if self.processing_thread.processor:
                self.processing_thread.processor = None
            self.processing_thread = None

        # Wait for organize thread to finish
        if hasattr(self, 'organize_thread') and self.organize_thread:
            if self.organize_thread.isRunning():
                self.organize_thread.request_stop()
                self.organize_thread.wait(3000)
            self.organize_thread = None

        # Wait for monitor thread to finish
        if hasattr(self, 'monitor_thread') and self.monitor_thread:
            if self.monitor_thread.isRunning():
                self.monitor_thread.wait(1000)
            self.monitor_thread = None

        # Release the processor (holds YOLO model and EasyOCR)
        if self.processor:
            self.processor = None

        # Force garbage collection before Qt cleanup
        gc.collect()

    # Slots
    @Slot(str, str)
    def _on_process_requested(self, input_dir: str, output_dir: str):
        """Handle process request from processing panel."""
        # This WIPES all in-memory review state (viewed/checked/cropped/note/
        # label) and reprocesses from scratch. If this fires unexpectedly
        # during review it looks exactly like "it reset / auto-archived".
        dlog("PROCESS_REQUESTED (clears results!)", input_dir=input_dir,
             discarding=fingerprint(self.current_results),
             was_processing=self.is_processing, monitoring=self.is_monitoring)
        self.is_processing = True
        self.status_label.setText(f"Processing: {input_dir}")
        self.processing_panel.set_processing(True)
        self.preview_panel.set_batch_processing_active(True)

        # Clear previous results when starting a new batch
        self.current_results = []
        self.results_list.clear()
        self._current_input_dir = input_dir
        self._session_dirty = False

        # Save directories to settings
        self.settings.ui.last_input_dir = input_dir
        self.settings.ui.last_output_dir = output_dir
        self.settings.save()

        # Start processing in background thread
        self._start_processing(input_dir, output_dir)

    @Slot(str)
    def _on_organize_requested(self, input_dir: str):
        """Handle organize request from processing panel."""
        from .processing_thread import OrganizeThread

        self.is_processing = True
        self.status_label.setText(f"Organizing: {input_dir}")
        self.processing_panel.set_processing(True)

        # Create and start organize thread
        self.organize_thread = OrganizeThread(
            input_dir,
            use_gpu=self.settings.processing.use_gpu
        )
        self.organize_thread.progress_updated.connect(self._on_organize_progress)
        self.organize_thread.organize_complete.connect(self._on_organize_complete)
        self.organize_thread.error_occurred.connect(self._on_organize_error)
        self.organize_thread.start()

    @Slot(int, int, str)
    def _on_organize_progress(self, current: int, total: int, message: str):
        """Handle organize progress update."""
        if total > 0:
            self.processing_panel.progress_bar.setMaximum(total)
            self.processing_panel.progress_bar.setValue(current)
        self.status_label.setText(message)

    @Slot(dict)
    def _on_organize_complete(self, result: dict):
        """Handle organize completion."""
        self.is_processing = False
        self.processing_panel.set_processing(False)
        self.processing_panel.progress_bar.setValue(0)

        pairs = result.get('pairs_organized', 0)
        corrected = result.get('images_corrected', 0)
        elapsed = result.get('time_taken', 0)

        self.status_label.setText(
            f"Organized {pairs} pairs ({corrected} corrected) in {elapsed:.1f}s - "
            f"Files renamed to Dollar_001.jpg format"
        )

        # Show info dialog
        from PySide6.QtWidgets import QMessageBox
        QMessageBox.information(
            self,
            "Organize Complete",
            f"Folder organized successfully!\n\n"
            f"Pairs organized: {pairs}\n"
            f"Images corrected: {corrected}\n"
            f"Time taken: {elapsed:.1f}s\n\n"
            f"Files renamed to Dollar_001.jpg through Dollar_{pairs*2:03d}.jpg\n\n"
            f"The folder is now ready for faster processing."
        )

    @Slot(str)
    def _on_organize_error(self, error: str):
        """Handle organize error."""
        self.is_processing = False
        self.processing_panel.set_processing(False)
        self.processing_panel.progress_bar.setValue(0)
        self.status_label.setText(f"Error: {error}")

        from PySide6.QtWidgets import QMessageBox
        QMessageBox.critical(self, "Organize Error", error)

    @Slot()
    def _on_stop_requested(self):
        """Handle stop request."""
        if self.processing_thread and self.processing_thread.isRunning():
            self.processing_thread.request_stop()
        self.is_processing = False
        self.status_label.setText("Stopping...")
        self.processing_panel.set_processing(False)
        self.preview_panel.set_batch_processing_active(False)

    @Slot(dict)
    def _on_result_selected(self, result: dict):
        """Handle selection of a result item."""
        self.preview_panel.show_bill(result)
        # Mark dirty so autosave captures viewed status
        self._mark_session_dirty()

    @Slot(str, str, str)
    def _on_correction_applied(self, filename: str, original: str, corrected: str):
        """Handle a correction applied via context menu."""
        # Save to correction manager
        self.correction_manager.add_correction(filename, original, corrected)
        self.correction_manager.save()

        # Update status
        self.status_label.setText(f"Correction saved: {original} → {corrected}")

        # Mark session as dirty for autosave (correction modifies session state)
        self._mark_session_dirty()

        # Refresh preview if showing this bill
        selected = self.results_list.get_selected_result()
        if selected and selected.get('front_file') == filename:
            self.preview_panel.show_bill(selected)

    def _on_open_folder(self):
        """Handle open folder action."""
        # Priority: default_working_dir (user setting) > last_input_dir (history) > home
        start_dir = (self.settings.ui.default_working_dir or
                     self.settings.ui.last_input_dir or
                     str(Path.home()))
        folder = QFileDialog.getExistingDirectory(
            self, "Select Scan Folder", start_dir
        )
        if folder:
            self.processing_panel.set_input_dir(folder)

    def _on_settings(self):
        """Open settings dialog."""
        from .settings_dialog import SettingsDialog
        dialog = SettingsDialog(self.settings, self)
        if dialog.exec():
            self.settings.save()
            self._apply_settings()

    def _on_pattern_manager(self):
        """Open pattern manager dialog."""
        from .pattern_dialog import PatternDialog
        dialog = PatternDialog(self)
        dialog.exec()

        # Reload preview panel's pattern engine if patterns were modified
        if dialog.patterns_were_modified():
            self.preview_panel.reload_pattern_engine()

    def _on_crop_manager(self):
        """Open eBay crop manager dialog."""
        import yaml
        from .crop_dialog import EbayCropDialog

        # Load current config
        config_path = Path(__file__).parent.parent / "config.yaml"
        if config_path.exists():
            with open(config_path) as f:
                config = yaml.safe_load(f)
        else:
            config = {}

        dialog = EbayCropDialog(config, self)
        if dialog.exec():
            # Save updated config
            updated_config = dialog.get_config()
            with open(config_path, 'w') as f:
                yaml.dump(updated_config, f, default_flow_style=False, sort_keys=False)
            QMessageBox.information(
                self, "Settings Saved",
                "Crop settings have been saved.\nChanges will apply to future processing."
            )

    def _toggle_review_filter(self, checked: bool):
        """Toggle showing only review items."""
        self.results_list.set_filter('needs_review', checked)

    def _toggle_fancy_filter(self, checked: bool):
        """Toggle showing only fancy bills."""
        self.results_list.set_filter('is_fancy', checked)

    def _toggle_serial_region(self, checked: bool):
        """Toggle serial region panel visibility."""
        self.preview_panel.set_serial_region_visible(checked)
        self.settings.ui.show_serial_region = checked
        self.settings.save()

    def _toggle_details(self, checked: bool):
        """Toggle bill details panel visibility."""
        self.preview_panel.set_details_visible(checked)
        self.settings.ui.show_bill_details = checked
        self.settings.save()

    def _on_layout_changed(self, layout_id: str):
        """Handle layout selection from menu."""
        # Update menu checkmarks
        for lid, action in self.layout_actions.items():
            action.setChecked(lid == layout_id)

        # Apply the layout
        self.layout_manager.apply_layout(layout_id)

        # Save preference
        self.settings.ui.layout_mode = layout_id
        self.settings.save()

    def _setup_columns_menu(self, menu: QMenu):
        """Populate the Columns submenu with toggle actions for each column."""
        columns = self.results_list.get_column_info()
        for col_idx, col_name, col_tooltip, visible in columns:
            action = QAction(col_name, self, checkable=True)
            action.setChecked(visible)
            if col_tooltip:
                action.setToolTip(col_tooltip)
            action.triggered.connect(
                lambda checked, idx=col_idx: self._toggle_column_visibility(idx, checked)
            )
            menu.addAction(action)
            self.column_actions[col_idx] = action

    def _toggle_column_visibility(self, column: int, visible: bool):
        """Toggle visibility of a results list column."""
        self.results_list.set_column_visible(column, visible)

    def _on_column_hidden_from_header(self, column: int, visible: bool):
        """Called when a column is hidden via header right-click menu.

        Updates the View > Columns menu to reflect the change.
        """
        if column in self.column_actions:
            self.column_actions[column].setChecked(visible)

    def _refresh_view(self):
        """Refresh the current view."""
        self.results_list.refresh()

    def _on_about(self):
        """Show about dialog."""
        QMessageBox.about(
            self, "About Dollar Bill Processor",
            f"Dollar Bill Processor v{get_version_string()}\n\n"
            "Automated detection of fancy serial numbers\n"
            "on US currency bills.\n\n"
            "Features:\n"
            "- YOLO-based serial number detection\n"
            "- 50+ pattern recognition rules\n"
            "- Manual correction workflow\n\n"
            "Built with PySide6 and OpenCV"
        )

    def _export_results(self, format_type: str):
        """Export results to file."""
        if not self.current_results:
            QMessageBox.warning(self, "No Results", "No results to export.")
            return

        if format_type == "csv":
            path, _ = QFileDialog.getSaveFileName(
                self, "Export CSV", "", "CSV Files (*.csv)"
            )
            if path:
                self._export_csv(path)
        elif format_type == "excel":
            path, _ = QFileDialog.getSaveFileName(
                self, "Export Excel", "", "Excel Files (*.xlsx)"
            )
            if path:
                self._export_excel(path)
        elif format_type == "html":
            path, _ = QFileDialog.getSaveFileName(
                self, "Export HTML", "", "HTML Files (*.html)"
            )
            if path:
                self._export_html(path)

    def _auto_export(self, summary: dict):
        """Auto-export CSV and/or summary if enabled in settings."""
        from datetime import datetime
        input_dir = Path(self.settings.ui.last_input_dir)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

        exported = []

        # Auto-export CSV
        if self.settings.export.auto_export_csv:
            csv_path = input_dir / f"results_{timestamp}.csv"
            self._export_csv(str(csv_path), quiet=True)
            exported.append(f"CSV: {csv_path.name}")

        # Auto-export summary
        if self.settings.export.auto_export_summary:
            summary_path = input_dir / f"summary_{timestamp}.txt"
            self._export_summary(str(summary_path), summary)
            exported.append(f"Summary: {summary_path.name}")

        if exported:
            self.status_label.setText(
                f"Complete: {summary.get('total', 0)} bills | Auto-exported: {', '.join(exported)}"
            )

    def _export_summary(self, path: str, summary: dict):
        """Export processing summary to text file."""
        from datetime import datetime
        fancy_bills = [r for r in self.current_results if r.get('is_fancy')]

        with open(path, 'w') as f:
            f.write("Dollar Bill Processing Summary\n")
            f.write("=" * 40 + "\n")
            f.write(f"Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
            f.write(f"Input: {self.settings.ui.last_input_dir}\n\n")

            f.write(f"Total bills processed: {summary.get('total', 0)}\n")
            f.write(f"Fancy bills found: {summary.get('fancy_count', 0)}\n")
            f.write(f"Bills needing review: {summary.get('review_count', 0)}\n\n")

            if fancy_bills:
                f.write("Fancy Bills:\n")
                f.write("-" * 40 + "\n")
                for bill in fancy_bills:
                    f.write(f"  {bill.get('serial', 'N/A')}: {bill.get('fancy_types', '')}\n")

    def _export_csv(self, path: str, quiet: bool = False):
        """Export results to CSV."""
        import csv
        with open(path, 'w', newline='') as f:
            writer = csv.DictWriter(f, fieldnames=[
                'position', 'front_file', 'back_file', 'serial', 'fancy_types',
                'confidence', 'baseline_variance', 'seal_x', 'seal_y', 'seal_containment',
                'is_fancy', 'needs_review', 'serial_region_path', 'error',
                'front_align_angle', 'front_align_flipped',
                'series_year', 'front_plate', 'back_plate', 'potential_mule', 'serial_mismatch',
                'viewed', 'cropped', 'sent_for_review', 'checked',
                'note', 'pattern_override', 'pattern_overrides'
            ], extrasaction='ignore')
            writer.writeheader()
            writer.writerows(self.current_results)
        if not quiet:
            self.status_label.setText(f"Exported to {path}")

    def _export_excel(self, path: str):
        """Export results to Excel."""
        try:
            import pandas as pd
            df = pd.DataFrame(self.current_results)
            df.to_excel(path, index=False)
            self.status_label.setText(f"Exported to {path}")
        except ImportError:
            QMessageBox.warning(
                self, "Missing Dependency",
                "pandas and openpyxl are required for Excel export.\n"
                "Install with: pip install pandas openpyxl"
            )

    def _export_html(self, path: str):
        """Export results to HTML report."""
        html = self._generate_html_report()
        with open(path, 'w') as f:
            f.write(html)
        self.status_label.setText(f"Exported to {path}")

    def _generate_html_report(self) -> str:
        """Generate HTML report content."""
        fancy_bills = [r for r in self.current_results if r.get('is_fancy')]
        review_bills = [r for r in self.current_results if r.get('needs_review')]

        html = f"""<!DOCTYPE html>
<html>
<head>
    <title>Dollar Bill Processing Report</title>
    <style>
        body {{ font-family: Arial, sans-serif; margin: 20px; }}
        h1 {{ color: #2e7d32; }}
        table {{ border-collapse: collapse; width: 100%; margin: 20px 0; }}
        th, td {{ border: 1px solid #ddd; padding: 8px; text-align: left; }}
        th {{ background-color: #4CAF50; color: white; }}
        tr:nth-child(even) {{ background-color: #f2f2f2; }}
        .fancy {{ background-color: #c8e6c9; }}
        .review {{ background-color: #fff3e0; }}
        .summary {{ background-color: #e3f2fd; padding: 15px; margin: 10px 0; }}
    </style>
</head>
<body>
    <h1>Dollar Bill Processing Report</h1>

    <div class="summary">
        <strong>Summary:</strong><br>
        Total Bills: {len(self.current_results)}<br>
        Fancy Bills: {len(fancy_bills)}<br>
        Needs Review: {len(review_bills)}
    </div>

    <h2>Fancy Bills</h2>
    <table>
        <tr><th>Position</th><th>Serial</th><th>Patterns</th><th>Confidence</th></tr>
"""
        for r in fancy_bills:
            html += f"""        <tr class="fancy">
            <td>{r.get('position', '')}</td>
            <td>{r.get('serial', '')}</td>
            <td>{r.get('fancy_types', '')}</td>
            <td>{r.get('confidence', '')}</td>
        </tr>
"""

        html += """    </table>

    <h2>All Bills</h2>
    <table>
        <tr><th>Position</th><th>File</th><th>Serial</th><th>Patterns</th><th>Status</th></tr>
"""
        for r in self.current_results:
            css_class = ""
            if r.get('is_fancy'):
                css_class = "fancy"
            elif r.get('needs_review'):
                css_class = "review"

            status = []
            if r.get('is_fancy'):
                status.append("Fancy")
            if r.get('needs_review'):
                status.append("Review")
            if r.get('error'):
                status.append(r['error'])

            html += f"""        <tr class="{css_class}">
            <td>{r.get('position', '')}</td>
            <td>{r.get('front_file', '')}</td>
            <td>{r.get('serial', '')}</td>
            <td>{r.get('fancy_types', '')}</td>
            <td>{', '.join(status) if status else 'OK'}</td>
        </tr>
"""

        html += """    </table>
</body>
</html>"""
        return html

    def _start_processing(self, input_dir: str, output_dir: str):
        """Start the processing in a background thread."""
        from .processing_thread import ProcessingThread

        self.processing_thread = ProcessingThread(
            input_dir=input_dir,
            output_dir=output_dir,
            use_gpu=self.settings.processing.use_gpu,
            verify_pairs=self.settings.processing.verify_pairs,
            crop_all=self.settings.processing.crop_all,
            auto_crop=self.settings.processing.auto_crop,
            extract_plate_info=self.settings.processing.extract_plate_info
        )
        self.processing_thread.progress_updated.connect(self._on_progress_updated)
        self.processing_thread.result_ready.connect(self._on_result_ready)
        self.processing_thread.processing_complete.connect(self._on_processing_complete)
        self.processing_thread.error_occurred.connect(self._on_processing_error)
        self.processing_thread.start()

    @Slot(int, int, str)
    def _on_progress_updated(self, current: int, total: int, message: str):
        """Handle progress update from processing thread."""
        self.processing_panel.update_progress(current, total)
        self.progress_label.setText(f"{current}/{total}")
        self.status_label.setText(message)

    @Slot(dict)
    def _on_result_ready(self, result: dict):
        """Handle a single result from processing."""
        self.current_results.append(result)
        self.results_list.add_result(result)
        # Mark session as dirty for autosave
        self._mark_session_dirty()
        # Force UI to update immediately (important for monitor mode)
        QApplication.processEvents()

    @Slot(dict)
    def _on_processing_complete(self, summary: dict):
        """Handle processing completion."""
        self.is_processing = False
        self.processing_panel.set_processing(False)
        self.preview_panel.set_batch_processing_active(False)

        # Grab the processor from the thread for alignment feature and preview panel
        if hasattr(self, 'processing_thread') and self.processing_thread:
            self.processor = self.processing_thread.processor
            # Share processor with preview panel to avoid re-loading YOLO
            self.preview_panel.set_processor(self.processor)

        total = summary.get('total', 0)
        fancy = summary.get('fancy_count', 0)
        review = summary.get('review_count', 0)

        self.status_label.setText(
            f"Complete: {total} bills processed, {fancy} fancy, {review} need review"
        )
        self.progress_label.setText("")

        # Trigger immediate autosave on batch completion
        self._trigger_autosave()

        # Auto-export if enabled
        if self.current_results:
            self._auto_export(summary)

        # Auto-archive if enabled, otherwise enable manual archive button
        dlog("processing.complete", total=total, fancy=fancy, review=review,
             auto_archive=self.settings.processing.auto_archive,
             state=fingerprint(self.current_results))
        if self.settings.processing.auto_archive and self.current_results:
            self._archive_manual_batch()

        # Update archive button state
        self.processing_panel.set_archive_available(
            available=bool(self.current_results),
            auto_archive_enabled=self.settings.processing.auto_archive
        )

    @Slot(str)
    def _on_processing_error(self, error: str):
        """Handle processing error."""
        self.is_processing = False
        self.processing_panel.set_processing(False)
        self.preview_panel.set_batch_processing_active(False)
        QMessageBox.critical(self, "Processing Error", error)
        self.status_label.setText(f"Error: {error}")

    def _apply_settings(self):
        """Apply changed settings to the UI."""
        # Apply theme and font size together (they share a stylesheet)
        self._apply_theme_and_font()

        # Update autosave timer in case interval changed
        self._update_autosave_timer()

        # Refresh batch list from archive directory
        self.results_list.refresh_batch_list()

        # Update archive button state based on new settings
        self.processing_panel.set_archive_available(
            available=bool(self.current_results),
            auto_archive_enabled=self.settings.processing.auto_archive
        )

    def _apply_theme_and_font(self):
        """Apply theme and font size to the application."""
        from .theme_manager import apply_theme, get_combined_stylesheet

        app = QApplication.instance()
        theme = self.settings.ui.theme
        font_size = self.settings.ui.font_size

        # Apply palette (handles most color changes)
        apply_theme(app, theme)

        # Apply combined stylesheet (font sizes + dark mode polish)
        stylesheet = get_combined_stylesheet(theme, font_size)
        app.setStyleSheet(stylesheet)

    # =========================================================================
    # Lazy Processor Creation
    # =========================================================================

    def _get_or_create_processor(self, silent: bool = False):
        """Get existing processor or create one lazily for cropping/alignment.

        This allows cropping from restored sessions without reprocessing.

        Args:
            silent: If True, don't ask user for confirmation (used for auto-load on restore)
        """
        if self.processor:
            return self.processor

        # Ask user if they want to load the processor (unless silent mode)
        if not silent:
            reply = QMessageBox.question(
                self, "Load Processor",
                "Cropping requires loading the YOLO model.\n\n"
                "This may take a few seconds. Continue?",
                QMessageBox.Yes | QMessageBox.No,
                QMessageBox.Yes
            )
            if reply != QMessageBox.Yes:
                return None

        # Show progress
        self.status_label.setText("Loading YOLO model...")
        QApplication.processEvents()

        try:
            from process_production import ProductionProcessor, Config

            # Find the YOLO model (resolve against the app base so it works both
            # from source and when frozen by PyInstaller). The ONNX backend loads
            # best.onnx derived from this path; best.pt only needs to exist.
            from resource_path import app_base
            base = app_base()
            model_path = base / "models" / "best.pt"
            if not model_path.exists():
                model_path = base / "best.pt"
            if not model_path.exists() and (base / "best.onnx").exists():
                model_path = base / "best.pt"  # ONNX sibling drives inference
            if not model_path.exists() and not (base / "best.onnx").exists():
                if not silent:
                    QMessageBox.warning(
                        self, "Model Not Found",
                        "Could not find YOLO model (best.onnx / best.pt).\n\n"
                        "Please ensure the model file is bundled with the app."
                    )
                self.status_label.setText("Ready")
                return None

            # Load config from config.yaml (for crop settings, etc.)
            script_dir = base
            config_path = script_dir / "config.yaml"
            patterns_dir = script_dir / "patterns"

            cfg = Config(config_path) if config_path.exists() else None

            # Create processor with config
            self.processor = ProductionProcessor(
                yolo_model_path=model_path,
                use_gpu=self.settings.processing.use_gpu,
                cfg=cfg,
                patterns_dir=patterns_dir if patterns_dir.exists() else None
            )

            # Share with preview panel
            self.preview_panel.set_processor(self.processor)

            self.status_label.setText("Ready - processor loaded")
            return self.processor

        except Exception as e:
            if not silent:
                QMessageBox.warning(
                    self, "Load Error",
                    f"Failed to load processor:\n\n{str(e)}"
                )
            else:
                print(f"[MainWindow] Failed to load processor: {e}")
            self.status_label.setText("Ready")
            return None

    def _load_processor_for_restored_session(self):
        """Load processor automatically after restoring a session."""
        if self.processor:
            return  # Already loaded

        if not self.current_results:
            return  # No session to work with

        print("[MainWindow] Auto-loading processor for restored session...")
        self._get_or_create_processor(silent=True)

    # =========================================================================
    # Autosave & Recovery Methods
    # =========================================================================

    def _setup_autosave_timer(self):
        """Setup the periodic autosave timer."""
        self._autosave_timer = QTimer(self)
        self._autosave_timer.timeout.connect(self._do_autosave)
        self._update_autosave_timer()

    def _update_autosave_timer(self):
        """Update autosave timer based on settings."""
        if self.settings.autosave.enabled:
            interval_ms = self.settings.autosave.interval_seconds * 1000
            self._autosave_timer.start(interval_ms)
        else:
            self._autosave_timer.stop()

    def _check_for_recovery(self):
        """Check for recovery file on startup and offer to restore."""
        if not self.recovery_manager.has_recovery_file():
            return

        info = self.recovery_manager.get_recovery_info()
        if not info or info.get("result_count", 0) == 0:
            # Empty or invalid recovery - clear it
            self.recovery_manager.clear_recovery()
            return

        # Show recovery dialog
        from .recovery_dialog import RecoveryDialog

        dialog = RecoveryDialog(info, self)
        dialog.exec()

        action = dialog.get_action()
        if action == RecoveryDialog.RESTORE:
            self._restore_session()
        elif action == RecoveryDialog.DISCARD:
            self.recovery_manager.clear_recovery()
            self.status_label.setText("Recovery discarded - starting fresh")

    def _restore_session(self):
        """Restore session from recovery file."""
        dlog("RESTORE_session.START", replacing=fingerprint(self.current_results))
        data = self.recovery_manager.load_recovery()
        if not data:
            QMessageBox.warning(self, "Recovery Error", "Failed to load recovery data.")
            return

        # Restore results
        results = data.get("results", [])

        # Normalize boolean fields (handle string "True"/"False" from corrupted data)
        for result in results:
            # Normalize is_fancy
            is_fancy = result.get("is_fancy", False)
            if isinstance(is_fancy, str):
                result["is_fancy"] = is_fancy.lower() == "true"

            # Normalize needs_review
            needs_review = result.get("needs_review", False)
            if isinstance(needs_review, str):
                result["needs_review"] = needs_review.lower() == "true"

            # Normalize front_align_flipped
            flipped = result.get("front_align_flipped", False)
            if isinstance(flipped, str):
                result["front_align_flipped"] = flipped.lower() == "true"

            # Normalize potential_mule
            mule = result.get("potential_mule", False)
            if isinstance(mule, str):
                result["potential_mule"] = mule.lower() == "true"

            # Normalize review status fields
            for key in ("viewed", "cropped", "sent_for_review", "checked"):
                val = result.get(key, False)
                if isinstance(val, str):
                    result[key] = val.lower() == "true"

            # Set alignment data flag if alignment angle is present
            # This enables the Align button to use cached values without needing YOLO
            if "front_align_angle" in result:
                result["_has_alignment_data"] = True

        self.current_results = results
        self._current_input_dir = data.get("input_directory", "")

        # Debug: log what we're restoring
        fancy_count = sum(1 for r in results if r.get("is_fancy"))
        review_count = sum(1 for r in results if r.get("needs_review"))
        has_align = sum(1 for r in results if r.get("front_align_angle", 0) != 0)
        print(f"[Recovery] Restoring {len(results)} results: {fancy_count} fancy, {review_count} needs_review, {has_align} with alignment")
        if results:
            r = results[0]
            print(f"[Recovery] Sample result: is_fancy={r.get('is_fancy')}, needs_review={r.get('needs_review')}, align_angle={r.get('front_align_angle')}")

        # Populate results list
        for result in results:
            self.results_list.add_result(result)

        # Update status
        count = len(results)
        complete = data.get("processing_complete", False)

        status = f"Restored: {count} bills"
        if fancy_count:
            status += f", {fancy_count} fancy"
        if not complete:
            status += " (processing was interrupted)"

        self.status_label.setText(status)

        # Set input directory in processing panel
        if self._current_input_dir:
            self.processing_panel.set_input_dir(self._current_input_dir)

        # Mark as clean (we just loaded from recovery)
        self._session_dirty = False

        # Enable archive button if we have results
        self.processing_panel.set_archive_available(
            available=bool(results),
            auto_archive_enabled=self.settings.processing.auto_archive
        )

        # Proactively load YOLO processor for cropping/alignment
        # Use a short delay to let the UI finish rendering first
        QTimer.singleShot(500, self._load_processor_for_restored_session)

    def _do_autosave(self):
        """Perform periodic autosave if there are changes."""
        if not self.settings.autosave.enabled:
            return

        # Only save if we have results and session is dirty
        if not self.current_results:
            return

        if not self._session_dirty:
            return

        # Normalize boolean fields before saving (prevent string "True"/"False" issues)
        for result in self.current_results:
            for key in ("is_fancy", "needs_review", "front_align_flipped", "potential_mule",
                        "viewed", "cropped", "sent_for_review", "checked"):
                val = result.get(key)
                if isinstance(val, str):
                    result[key] = val.lower() == "true"

        # Determine if processing is complete
        processing_complete = not self.is_processing and not self.is_monitoring

        # Get current selection index
        selected = self.results_list.tree.currentItem()
        last_index = -1
        if selected:
            last_index = self.results_list.tree.indexOfTopLevelItem(selected)

        # Save session
        success = self.recovery_manager.save_session(
            results=self.current_results,
            input_directory=self._current_input_dir,
            processing_complete=processing_complete,
            total_processed=len(self.current_results),
            last_selected_index=last_index
        )

        dlog("autosave", success=success, state=fingerprint(self.current_results))
        if success:
            self._session_dirty = False

    def _mark_session_dirty(self):
        """Mark the session as having unsaved changes."""
        self._session_dirty = True

    def _trigger_autosave(self):
        """Force an immediate autosave (e.g., after batch complete)."""
        if self.settings.autosave.enabled and self.current_results:
            self._session_dirty = True
            self._do_autosave()

    def _clear_recovery_after_archive(self):
        """Clear recovery file after successful archive."""
        self.recovery_manager.clear_recovery()
        self._session_dirty = False

    # =========================================================================
    # Monitor Mode Methods
    # =========================================================================

    def _on_monitor_mode_changed(self, enabled: bool):
        """Handle monitor mode checkbox toggle."""
        if enabled:
            # Update display with configured directories
            watch_dir = self.settings.monitor.watch_directory
            output_dir = self.settings.monitor.output_directory
            self.processing_panel.set_monitor_dirs(watch_dir, output_dir)

    def _start_monitoring(self):
        """Start monitor mode."""
        from .file_watcher import FileWatcher
        from .monitor_thread import MonitorThread

        # Validate settings
        watch_dir = self.settings.monitor.watch_directory
        output_dir = self.settings.monitor.output_directory

        if not watch_dir:
            QMessageBox.warning(
                self, "Configuration Required",
                "Please configure the watch directory in Settings > Monitor."
            )
            return

        # Expand user path (handle ~ on all platforms)
        watch_path = Path(watch_dir).expanduser().resolve()
        print(f"[MainWindow] Monitor watch path: {watch_path}")

        if not watch_path.exists():
            QMessageBox.warning(
                self, "Directory Not Found",
                f"Watch directory does not exist:\n{watch_dir}\n\n"
                "Please create the directory or configure a different path."
            )
            return

        if not output_dir:
            output_dir = str(watch_path / "fancy_bills")
        else:
            output_dir = str(Path(output_dir).expanduser().resolve())

        print(f"[MainWindow] Monitor output path: {output_dir}")

        # Switch to current session and clear previous results
        dlog("MONITOR_START (clears results!)",
             discarding=fingerprint(self.current_results))
        self.results_list.select_current_session()
        self.current_results = []
        self.results_list.clear()
        self.preview_panel.clear()
        self._current_input_dir = str(watch_path)
        self._session_dirty = False

        # Create monitor thread
        self.monitor_thread = MonitorThread(
            watch_dir=watch_path,
            output_dir=Path(output_dir),
            use_gpu=self.settings.processing.use_gpu,
            verify_pairs=self.settings.processing.verify_pairs,
            crop_all=self.settings.processing.crop_all,
            extract_plate_info=self.settings.processing.extract_plate_info
        )

        # Connect signals
        self.monitor_thread.progress_updated.connect(self._on_progress_updated)
        self.monitor_thread.result_ready.connect(self._on_result_ready)
        self.monitor_thread.processing_complete.connect(self._on_monitor_complete)
        self.monitor_thread.error_occurred.connect(self._on_processing_error)
        self.monitor_thread.status_updated.connect(self._on_monitor_status)

        # Create file watcher
        self.file_watcher = FileWatcher(
            watch_dir=watch_path,
            poll_interval=self.settings.monitor.poll_interval,
            settle_time=self.settings.monitor.file_settle_time
        )

        # Connect file watcher to monitor thread
        self.file_watcher.new_file_detected.connect(self.monitor_thread.handle_new_file)
        self.file_watcher.error_occurred.connect(self._on_processing_error)

        # Start threads
        print("[MainWindow] Starting monitor thread...")
        self.monitor_thread.start()
        print("[MainWindow] Starting file watcher...")
        self.file_watcher.start()

        print("[MainWindow] Monitor mode started successfully")

        # Update UI state
        self.is_monitoring = True
        self.processing_panel.set_monitoring(True)
        self.status_label.setText(f"Monitoring: {watch_dir}")

    def _stop_monitoring(self):
        """Stop monitor mode and optionally archive files."""
        dlog("MONITOR_STOP", is_monitoring=self.is_monitoring,
             auto_archive=self.settings.monitor.auto_archive,
             state=fingerprint(self.current_results))
        if not self.is_monitoring:
            return

        # Stop the file watcher
        if self.file_watcher:
            self.file_watcher.stop()
            self.file_watcher.wait(2000)
            self.file_watcher = None

        # Stop the monitor thread
        if self.monitor_thread:
            self.monitor_thread.stop()
            self.monitor_thread.wait(5000)

            # Grab processor for alignment feature
            self.processor = self.monitor_thread.processor

            # Archive if enabled
            print(f"[MainWindow] Auto-archive enabled: {self.settings.monitor.auto_archive}, pairs: {self.monitor_thread.pair_count}")
            if self.settings.monitor.auto_archive and self.monitor_thread.pair_count > 0:
                self._archive_batch()

            self.monitor_thread = None

        # Update UI state
        self.is_monitoring = False
        self.processing_panel.set_monitoring(False)
        self.status_label.setText("Monitoring stopped")

        # Refresh batch list to show newly archived batch
        self.results_list.refresh_batch_list()

    @Slot(str)
    def _on_monitor_status(self, message: str):
        """Handle status updates from monitor thread."""
        self.status_label.setText(message)

    @Slot(str)
    def _on_batch_changed(self, batch_path: str):
        """Handle batch selection change in results list."""
        print(f"[MainWindow] _on_batch_changed: batch_path='{batch_path}'")
        if batch_path:
            # Viewing archived batch - set input directory to allow reprocessing
            self.preview_panel.clear()
            self.status_label.setText(f"Viewing archived batch: {Path(batch_path).name}")
            # Set input directory to archived batch path for reprocessing
            print(f"[MainWindow] Setting input_dir to: {batch_path}")
            self.processing_panel.set_input_dir(batch_path)
        else:
            # Back to current session. Selecting an archive had replaced the
            # results display with the archive's dicts; without re-installing
            # the live results here, the list stays stuck on the archive even
            # though current_results (still autosaved) holds the live session.
            # Re-sharing current_results also restores object identity so edits
            # flow back to the authoritative list again.
            self.preview_panel.clear()
            self.results_list.set_results(self.current_results)
            self.status_label.setText("Current session")
            dlog("session.restored_to_list", state=fingerprint(self.current_results))

    @Slot(dict)
    def _on_monitor_complete(self, summary: dict):
        """Handle monitor mode completion."""
        total = summary.get('total', 0)
        fancy = summary.get('fancy_count', 0)
        review = summary.get('review_count', 0)
        pending_fronts = summary.get('pending_fronts', 0)
        pending_backs = summary.get('pending_backs', 0)

        status = f"Monitoring stopped: {total} pairs processed, {fancy} fancy"
        if pending_fronts or pending_backs:
            status += f" ({pending_fronts + pending_backs} unpaired files)"

        self.status_label.setText(status)
        self.progress_label.setText("")

        # Trigger immediate autosave on monitor completion
        self._trigger_autosave()

        # Auto-export if enabled and we have results
        if self.current_results:
            self._auto_export(summary)

    def _archive_batch(self):
        """Move (or copy) processed files to a timestamped archive directory."""
        import shutil
        from datetime import datetime

        dlog("ARCHIVE_monitor.START", has_thread=bool(self.monitor_thread),
             state=fingerprint(self.current_results))
        if not self.monitor_thread:
            return

        archive_base = self.settings.monitor.archive_directory
        if not archive_base:
            archive_base = str(Path(self.settings.monitor.watch_directory) / "archive")

        archive_path = Path(archive_base)
        archive_path.mkdir(parents=True, exist_ok=True)

        # Create timestamped batch directory
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        batch_dir = archive_path / f"batch_{timestamp}"
        batch_dir.mkdir(parents=True, exist_ok=True)

        # Use copy or move based on settings
        copy_mode = self.settings.processing.archive_copy_mode
        file_op = shutil.copy2 if copy_mode else shutil.move
        op_name = "copying" if copy_mode else "moving"

        # Move/copy all session files (processed + unpaired)
        all_files = self.monitor_thread.get_all_session_files()
        print(f"[MainWindow] Archiving ({op_name}) {len(all_files)} files to {batch_dir}")
        moved_count = 0
        path_mapping = {}  # old_path -> new_path

        for file_path in all_files:
            if file_path.exists():
                try:
                    dest = batch_dir / file_path.name
                    file_op(str(file_path), str(dest))
                    path_mapping[str(file_path)] = str(dest)
                    moved_count += 1
                except Exception as e:
                    print(f"[MainWindow] Error {op_name} {file_path.name}: {e}")
                    self.status_label.setText(f"Error {op_name} {file_path.name}: {e}")
            else:
                print(f"[MainWindow] File no longer exists: {file_path}")

        # Move/copy fancy_bills output to batch archive
        output_dir = Path(self.settings.monitor.output_directory or
                         (Path(self.settings.monitor.watch_directory) / "fancy_bills")).expanduser().resolve()

        fancy_moved = 0
        if output_dir.exists():
            fancy_items = list(output_dir.glob("*"))
            if fancy_items:
                # Create fancy_bills subfolder in batch archive
                batch_fancy_dir = batch_dir / "fancy_bills"
                batch_fancy_dir.mkdir(parents=True, exist_ok=True)

                for item_path in fancy_items:
                    try:
                        dest = batch_fancy_dir / item_path.name
                        if item_path.is_dir():
                            if copy_mode:
                                shutil.copytree(str(item_path), str(dest))
                            else:
                                shutil.move(str(item_path), str(dest))
                        else:
                            file_op(str(item_path), str(dest))
                        fancy_moved += 1
                    except Exception as e:
                        print(f"[MainWindow] Error {op_name} {item_path.name}: {e}")

                print(f"[MainWindow] Archived ({op_name}) {fancy_moved} items (files/folders) to {batch_fancy_dir}")

        # Update result paths to point to new archive locations
        for result in self.current_results:
            front_file = result.get('front_file', '')
            back_file = result.get('back_file', '')
            if front_file and front_file in path_mapping:
                result['front_file'] = path_mapping[front_file]
            if back_file and back_file in path_mapping:
                result['back_file'] = path_mapping[back_file]

        # Update the results list with new paths
        self.results_list.update_result_paths(path_mapping)

        # Export batch CSV (with updated paths)
        if self.current_results:
            csv_path = batch_dir / "results.csv"
            self._export_batch_csv(csv_path)

        action = "Copied" if copy_mode else "Archived"
        self.status_label.setText(
            f"{action} {moved_count} files + {fancy_moved} fancy crops to {batch_dir.name}"
        )

        # Clear recovery file after successful archive
        self._clear_recovery_after_archive()

        return batch_dir

    def _archive_manual_batch(self):
        """Move (or copy) processed files to a timestamped archive directory (for manual processing)."""
        import shutil
        from datetime import datetime

        dlog("ARCHIVE_manual.START", has_thread=bool(self.processing_thread),
             copy_mode=self.settings.processing.archive_copy_mode,
             state=fingerprint(self.current_results))
        if not self.processing_thread:
            return

        input_dir = self.processing_thread.input_dir
        output_dir = self.processing_thread.output_dir

        # Use monitor archive directory, or create one based on input dir
        archive_base = self.settings.monitor.archive_directory
        if not archive_base:
            archive_base = str(input_dir.parent / "archive")

        archive_path = Path(archive_base)
        archive_path.mkdir(parents=True, exist_ok=True)

        # Create timestamped batch directory
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        batch_dir = archive_path / f"batch_{timestamp}"
        batch_dir.mkdir(parents=True, exist_ok=True)

        # Use copy or move based on settings
        copy_mode = self.settings.processing.archive_copy_mode
        file_op = shutil.copy2 if copy_mode else shutil.move
        op_name = "copying" if copy_mode else "moving"

        # Get list of processed files from results
        files_to_archive = set()
        for result in self.current_results:
            front_file = result.get('front_file', '')
            back_file = result.get('back_file', '')
            if front_file:
                files_to_archive.add(Path(front_file))
            if back_file:
                files_to_archive.add(Path(back_file))

        # Move/copy all source files and track old->new path mapping
        moved_count = 0
        path_mapping = {}  # old_path -> new_path
        for file_path in files_to_archive:
            if file_path.exists():
                try:
                    dest = batch_dir / file_path.name
                    file_op(str(file_path), str(dest))
                    path_mapping[str(file_path)] = str(dest)
                    moved_count += 1
                except Exception as e:
                    print(f"[MainWindow] Error {op_name} {file_path.name}: {e}")

        # Move/copy fancy_bills output to batch archive
        fancy_moved = 0
        if output_dir.exists():
            fancy_items = list(output_dir.glob("*"))
            if fancy_items:
                # Create fancy_bills subfolder in batch archive
                batch_fancy_dir = batch_dir / "fancy_bills"
                batch_fancy_dir.mkdir(parents=True, exist_ok=True)

                for item_path in fancy_items:
                    try:
                        dest = batch_fancy_dir / item_path.name
                        if item_path.is_dir():
                            if copy_mode:
                                shutil.copytree(str(item_path), str(dest))
                            else:
                                shutil.move(str(item_path), str(dest))
                        else:
                            file_op(str(item_path), str(dest))
                        fancy_moved += 1
                    except Exception as e:
                        print(f"[MainWindow] Error {op_name} {item_path.name}: {e}")

        # Update result paths to point to new archive locations
        for result in self.current_results:
            front_file = result.get('front_file', '')
            back_file = result.get('back_file', '')
            if front_file and front_file in path_mapping:
                result['front_file'] = path_mapping[front_file]
            if back_file and back_file in path_mapping:
                result['back_file'] = path_mapping[back_file]

        # Update the results list with new paths
        self.results_list.update_result_paths(path_mapping)

        # Export batch CSV (with updated paths)
        if self.current_results:
            csv_path = batch_dir / "results.csv"
            self._export_batch_csv(csv_path)

        action = "Copied" if copy_mode else "Archived"
        self.status_label.setText(
            f"{action} {moved_count} files + {fancy_moved} fancy crops to {batch_dir.name}"
        )

        # Clear recovery file after successful archive
        self._clear_recovery_after_archive()

        # Refresh batch list to show newly archived batch
        self.results_list.refresh_batch_list()

        return batch_dir

    def _on_archive_requested(self):
        """Handle manual archive button click."""
        if not self.current_results:
            return

        self._archive_manual_batch()
        self.processing_panel.reset_archive_button()

    def _export_batch_csv(self, csv_path: Path):
        """Export results to batch CSV file."""
        import csv
        with open(csv_path, 'w', newline='') as f:
            writer = csv.DictWriter(f, fieldnames=[
                'position', 'front_file', 'back_file', 'serial', 'fancy_types',
                'confidence', 'baseline_variance', 'seal_x', 'seal_y', 'seal_containment',
                'is_fancy', 'needs_review', 'serial_region_path', 'error',
                'front_align_angle', 'front_align_flipped',
                'series_year', 'front_plate', 'back_plate', 'potential_mule', 'serial_mismatch',
                'viewed', 'cropped', 'sent_for_review', 'checked',
                'note', 'pattern_override', 'pattern_overrides'
            ], extrasaction='ignore')
            writer.writeheader()
            writer.writerows(self.current_results)


def run_gui():
    """Launch the GUI application."""
    from debug_logger import dlog, get_log_path
    dlog("app.start", log=get_log_path())
    print(f"[DollarBill] Debug log: {get_log_path()}")

    app = QApplication(sys.argv)
    app.setApplicationName("Dollar Bill Processor")
    app.setOrganizationName("DollarBillProcessor")

    # Apply theme and font size from settings
    from .theme_manager import apply_theme, get_combined_stylesheet
    settings = get_settings()
    apply_theme(app, settings.ui.theme)
    stylesheet = get_combined_stylesheet(settings.ui.theme, settings.ui.font_size)
    app.setStyleSheet(stylesheet)

    window = MainWindow()
    window.show()

    sys.exit(app.exec())


if __name__ == "__main__":
    run_gui()
